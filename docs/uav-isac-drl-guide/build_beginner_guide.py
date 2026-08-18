#!/usr/bin/env python3
"""Build the beginner-friendly UAV--ISAC DRL guide as a Word document."""

from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
HERE = Path(__file__).resolve().parent
OUTPUT = HERE / "UAV_ISAC_DRL_Beginners_Guide.docx"

NAVY = "17365D"
BLUE = "2E75B6"
LIGHT_BLUE = "DDEBF7"
GREEN = "E2F0D9"
AMBER = "FFF2CC"
RED = "FCE4D6"
GREY = "E7E6E6"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Right-click and select Update Field to build the table of contents."
    separate.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, end])


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: list[str], numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        doc.add_paragraph(item, style=style)


def add_callout(doc: Document, title: str, text: str, color: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, color)
    p = cell.paragraphs[0]
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run(text)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        shade(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2:
                shade(cells[index], "F7F9FC")
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    doc.add_paragraph()


def add_code(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, "F2F2F2")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    doc.add_paragraph()


def make_feedback_figure(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(11, 4.8))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 5)
    ax.axis("off")
    boxes = [
        (0.3, 1.7, 2.0, 1.3, "ns-3 world", "UAVs, cows, NR, energy", "#DDEBF7"),
        (3.1, 1.7, 2.0, 1.3, "Observation", "50 numbers every 10 s", "#E2F0D9"),
        (5.9, 1.7, 2.0, 1.3, "Double DQN", "estimates 25 action values", "#FFF2CC"),
        (8.7, 1.7, 2.0, 1.3, "Live action", "sense, speed, altitude", "#FCE4D6"),
    ]
    for x, y, w, h, title, subtitle, color in boxes:
        patch = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.04",
                               linewidth=1.5, edgecolor="#17365D", facecolor=color)
        ax.add_patch(patch)
        ax.text(x + w / 2, y + 0.82, title, ha="center", va="center",
                fontsize=13, fontweight="bold", color="#17365D")
        ax.text(x + w / 2, y + 0.38, subtitle, ha="center", va="center", fontsize=9)
    for x1, x2 in [(2.3, 3.1), (5.1, 5.9), (7.9, 8.7)]:
        ax.add_patch(FancyArrowPatch((x1, 2.35), (x2, 2.35), arrowstyle="-|>",
                                     mutation_scale=16, linewidth=1.8, color="#2E75B6"))
    ax.add_patch(FancyArrowPatch((9.7, 1.65), (1.3, 1.65), connectionstyle="arc3,rad=-0.33",
                                 arrowstyle="-|>", mutation_scale=16, linewidth=1.8,
                                 color="#2E75B6"))
    ax.text(5.5, 0.35, "The action changes the same running episode; the next window reveals the consequence.",
            ha="center", fontsize=11, color="#17365D")
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def make_split_figure(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(10, 3.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 3.5)
    ax.axis("off")
    items = [
        (0.4, 6.0, "TRAIN", "Learn network weights\nSeeds 3100000–3100999", "#DDEBF7"),
        (6.6, 1.3, "VALIDATE", "Choose settings/checkpoint\nSeeds 7001–7010", "#FFF2CC"),
        (8.2, 1.3, "TEST", "One final comparison\nSeeds 9001–9010", "#E2F0D9"),
    ]
    for x, w, title, subtitle, color in items:
        patch = FancyBboxPatch((x, 1.0), w, 1.5, boxstyle="round,pad=0.04",
                               linewidth=1.5, edgecolor="#17365D", facecolor=color)
        ax.add_patch(patch)
        ax.text(x + w / 2, 1.95, title, ha="center", fontsize=13, fontweight="bold")
        ax.text(x + w / 2, 1.35, subtitle, ha="center", fontsize=9)
    ax.text(5, 0.25, "Never choose a model after looking at test performance.",
            ha="center", fontsize=11, color="#C00000", fontweight="bold")
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color in [("Title", 26, NAVY), ("Heading 1", 18, NAVY),
                              ("Heading 2", 14, BLUE), ("Heading 3", 11, NAVY)]:
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
    if "Keyword" not in styles:
        style = styles.add_style("Keyword", WD_STYLE_TYPE.CHARACTER)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE)
    header = section.header.paragraphs[0]
    header.text = "UAV–ISAC Deep Reinforcement Learning | Beginner’s Guide"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor.from_string("777777")
    add_page_number(section.footer.paragraphs[0])


def main() -> None:
    HERE.mkdir(parents=True, exist_ok=True)
    feedback = HERE / "drl_feedback_loop.png"
    split = HERE / "train_validation_test.png"
    make_feedback_figure(feedback)
    make_split_figure(split)

    doc = Document()
    configure_document(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("UAV-Enabled ISAC for Livestock Monitoring").bold = True
    p.runs[0].font.size = Pt(28)
    p.runs[0].font.color.rgb = RGBColor.from_string(NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("A from-zero guide to the genuine Deep Reinforcement Learning implementation")
    r.font.size = Pt(17)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("What the system does • what every important word means • what was changed • "
              "how it was tested • how to run it • how to report it honestly")
    doc.add_paragraph()
    add_callout(
        doc,
        "The one-sentence version",
        "We turned a simulator that previously tested one fixed UAV configuration per run "
        "into an interactive learning environment: every 10 simulated seconds, a Double-DQN "
        "agent observes the current mission, chooses one UAV control action, and sees the "
        "consequence in the next 10-second window.",
        GREEN,
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Prepared from the verified code and artifacts on 18 August 2026").italic = True
    doc.add_page_break()

    add_heading(doc, "How to use this guide", 1)
    add_body(doc, "This document is intentionally written for a reader who has never studied reinforcement learning, ns-3, ISAC, or UAV control. Read Chapters 1–5 for the story and vocabulary. Chapters 6–13 explain the technical design slowly. Chapters 14–18 show testing, results, commands, paper-writing rules, and remaining work.")
    add_callout(doc, "Important scientific warning", "The completed three-episode experiment is a smoke test. It proves that the software loop runs; it does not prove that DRL is better. A full multi-seed campaign is still required before adding a DRL improvement claim to the conference paper.", RED)
    add_heading(doc, "Contents", 1)
    add_toc(doc.add_paragraph())
    doc.add_page_break()

    add_heading(doc, "1. The problem in everyday language", 1)
    add_body(doc, "Imagine a large farm with 80 cows and four flying robots. A farmer wants to know where the cows are. The drones must fly around, detect cows, estimate their positions, and send reports through a 5G base station. They cannot waste battery, and the radio link is not perfect.")
    add_body(doc, "The difficult part is that a good decision changes over time. When few cows are detected, a drone may need to sense more frequently or move differently. When battery use is high, it may need to reduce effort. A fixed setting selected before take-off cannot react to what happens during the mission. DRL is introduced to learn these repeated decisions.")
    add_bullets(doc, [
        "Sensing question: Are the cows detected, and how accurate are their estimated locations?",
        "Communication question: Do reports arrive quickly, without loss, and at useful throughput?",
        "Mobility question: Where should each drone fly, how fast, and at what altitude?",
        "Energy question: Can the mission be completed within each drone’s battery budget?",
        "Learning question: Which action now is likely to produce good total performance later?",
    ])

    add_heading(doc, "2. The complete system at a glance", 1)
    doc.add_picture(str(feedback), width=Inches(6.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_body(doc, "One loop works as follows: ns-3 simulates ten seconds; it summarizes the new sensing, network, energy, position, and herd information; Python receives 50 numbers; Double DQN chooses one of 25 actions; ns-3 changes the live UAV object; then another ten seconds are simulated. This repeats until time expires or a battery is exhausted.")
    add_callout(doc, "Why the word live matters", "Changing only a Python variable or a C++ command-line value after a simulation starts does not control the current drone. The implemented setters modify the mobility and sensing objects that are already running inside ns-3.")

    add_heading(doc, "3. Essential vocabulary—from the beginning", 1)
    glossary = [
        ["UAV", "Unmanned Aerial Vehicle", "A drone. Four are simulated with different roles."],
        ["ISAC", "Integrated Sensing and Communication", "The UAV senses cows and communicates reports in one system."],
        ["5G NR", "Fifth-generation New Radio", "The cellular radio technology used for UAV-to-network traffic."],
        ["gNB", "Next-generation Node B", "The 5G base station."],
        ["UE", "User Equipment", "A device connected to the gNB; each UAV acts as a UE."],
        ["ns-3", "Network Simulator 3", "A discrete-event simulator that models time, packets, nodes, mobility, and radio links."],
        ["KPI", "Key Performance Indicator", "A measured quantity such as detection probability, delay, or energy."],
        ["RL", "Reinforcement Learning", "Learning actions by interacting with an environment and receiving rewards."],
        ["DRL", "Deep Reinforcement Learning", "RL that uses a neural network to represent a policy or value function."],
        ["Agent", "Decision maker", "Here, the Double-DQN program that chooses UAV actions."],
        ["Environment", "The world the agent controls", "Here, the running ns-3 livestock/UAV simulation."],
        ["State/observation", "What the agent sees", "A 50-number summary at a decision time."],
        ["Action", "What the agent does", "One of 25 sensing, speed, altitude, or no-op choices."],
        ["Reward", "Immediate score", "A number rewarding sensing/network quality and penalising energy."],
        ["Policy", "Decision rule", "A mapping from an observed state to an action."],
        ["Episode", "One complete mission", "One ns-3 run from reset to battery termination or time limit."],
        ["Step", "One decision cycle", "One action followed by one 10-second simulator window."],
        ["Transition", "One learning example", "State, action, reward, next state, and done flag."],
        ["DQN", "Deep Q-Network", "A neural network that estimates the long-term value of every discrete action."],
        ["Double DQN", "Two-network target rule", "Online network selects the next action; target network evaluates it."],
        ["Q-value", "Expected discounted return", "How good an action is now, including likely future rewards."],
        ["Bellman target", "Learning target", "Immediate reward plus discounted value of the next state."],
        ["Replay memory", "Transition storage", "A buffer that mixes old experiences into training batches."],
        ["Target network", "Slow reference network", "A copied network that stabilises the Bellman targets."],
        ["Epsilon", "Exploration probability", "Chance of taking a random action instead of the current greedy choice."],
        ["Hyperparameter", "A chosen learning setting", "For example learning rate, batch size, or discount factor."],
        ["Seed", "Random-number starting point", "Makes stochastic experiments repeatable and enables matched comparisons."],
        ["Baseline", "Reference method", "Static, random, rule-based, or RF control used for comparison."],
        ["RF", "Random Forest", "A supervised regression benchmark; it is not RL in this project."],
        ["O-RAN/xApp", "Programmable RAN controller", "A possible deployment route for a trained, safety-checked policy."],
        ["Smoke test", "Small execution check", "Confirms the pipeline runs; it is not a performance experiment."],
    ]
    add_table(doc, ["Keyword", "Plain expansion", "Meaning here"], glossary, [1.2, 1.8, 3.6])

    add_heading(doc, "4. What was wrong with the earlier ‘RL’ claim?", 1)
    add_body(doc, "The released file called offline_rl_analysis.py trained a Random Forest regressor. Each row described a complete fixed simulator configuration and its final J score. The model predicted which new fixed configurations might score well. That is supervised surrogate modelling plus candidate search.")
    add_table(doc, ["Question", "Released RF method", "Implemented DRL method"], [
        ["Does it act during an episode?", "No", "Yes, every 10 seconds"],
        ["Does it use ordered state transitions?", "No", "Yes"],
        ["Does an action affect the next state?", "No sequential link", "Yes, through live ns-3 control"],
        ["Bellman equation?", "No", "Yes"],
        ["Replay memory?", "No", "Yes"],
        ["Target neural network?", "No", "Yes"],
        ["Correct name", "RF surrogate/random-search benchmark", "Double-DQN controller"],
    ])
    add_callout(doc, "This does not make the RF method useless", "The RF method remains a valuable baseline. It answers a different question: which fixed configuration looks promising? The DRL agent answers: given the mission now, what should I change for the next window?", GREEN)

    add_heading(doc, "5. Reinforcement learning with a simple analogy", 1)
    add_body(doc, "Think of teaching someone to manage four sprinklers on a farm. Every ten minutes they see soil moisture, weather, and water remaining. They can adjust one sprinkler. Healthy crops give positive points; wasted water gives negative points. After many seasons, they learn which adjustments tend to produce good long-term outcomes. Replace sprinklers with UAV controls, soil information with KPIs, and seasons with simulated missions: that is the basic RL idea.")
    add_bullets(doc, [
        "The agent is not given a perfect answer for every state.",
        "It tries actions, sees consequences, and stores experiences.",
        "It must explore enough to discover alternatives.",
        "It must also exploit what it has already learned.",
        "A reward must represent the actual scientific objective; otherwise the agent learns the wrong behaviour very efficiently.",
    ])

    add_heading(doc, "6. What exactly is one episode and one step?", 1)
    add_body(doc, "An episode is one complete ns-3 mission. At reset, cow motion, radio randomness, and UAV motion are initialised using a recorded seed. With the standard 300-second mission and a 10-second control interval, there are approximately 29 actions after the first observation at 10 seconds. The final transition arrives at 300 seconds.")
    add_table(doc, ["Simulated time", "What happens"], [
        ["0–10 s", "ns-3 runs without an RL action to produce the first meaningful window."],
        ["10 s", "Python receives state 0 and sends action 0."],
        ["10–20 s", "The selected action affects live sensing or mobility."],
        ["20 s", "Python receives reward and state 1, then sends the next action."],
        ["…", "The locked state/action sequence continues."],
        ["300 s", "The time-limit transition is marked truncated; the episode ends."],
        ["Battery exhausted", "The transition is marked terminated; the episode ends early."],
    ])

    add_heading(doc, "7. The 50-number observation", 1)
    add_body(doc, "A neural network cannot directly read ‘the patrol UAV is high and the link is slow.’ It receives numbers. The observation deliberately combines what the agent needs to understand quality, physical state, active controls, and remaining mission time.")
    add_table(doc, ["Indices", "Count", "Contents", "Why included"], [
        ["0–4", "5", "Detection, RMSE, throughput, delay, loss", "Quality of the most recent window"],
        ["5–8", "4", "Remaining battery fraction per UAV", "Avoid energy exhaustion"],
        ["9–32", "24", "x, y, z, vx, vy, vz per UAV", "Know fleet position and motion"],
        ["33–36", "4", "Herd centroid and x/y spread", "Compact description of cow distribution"],
        ["37–40", "4", "Current sensing intervals", "Know the control already active"],
        ["41–44", "4", "Current target speeds", "Know live mobility targets"],
        ["45–48", "4", "Current target altitudes", "Know live altitude targets"],
        ["49", "1", "Normalised mission time", "Distinguish early and late decisions"],
    ])
    add_callout(doc, "Normalisation", "Raw units are placed on comparable scales. Horizontal coordinates and velocities can be negative and use −1 to 1. Most other values use 0 to 1. Normalisation helps neural-network optimisation and the environment rejects values outside its declared bounds.")

    add_heading(doc, "8. The 25 possible actions", 1)
    actions = [["0", "All", "No operation", "Keep current settings"]]
    roles = ["Surveillance", "Patrol", "Rapid-response", "Strategic"]
    operations = [
        ("Decrease sensing interval", "−0.1 s"),
        ("Increase sensing interval", "+0.1 s"),
        ("Decrease target speed", "−1 m/s"),
        ("Increase target speed", "+1 m/s"),
        ("Decrease target altitude", "−5 m"),
        ("Increase target altitude", "+5 m"),
    ]
    action_id = 1
    for role in roles:
        for operation, amount in operations:
            actions.append([str(action_id), role, operation, amount])
            action_id += 1
    add_table(doc, ["ID", "UAV role", "Operation", "Change"], actions)
    add_body(doc, "Every action is clipped to a safe role-specific bound. For example, repeated altitude-increase actions stop at the allowed upper altitude rather than increasing forever. Action 0 is important because sometimes changing nothing is best.")

    add_heading(doc, "9. The reward: how the agent is told what ‘good’ means", 1)
    add_body(doc, "The immediate window score rewards detection and throughput, penalises localisation error, delay, and loss, and then subtracts an energy term:")
    add_code(doc, "J = 3 × detection + 3 × throughput − 2 × RMSE − delay − loss\nreward = J − energy_weight × (window_energy / window_energy_budget)")
    add_body(doc, "Every metric in J is normalised first. The weights 3, 3, 2, 1, and 1 express priorities. They are design decisions, not laws of nature, and must be justified or tested with sensitivity analysis.")
    add_heading(doc, "9.1 A small made-up numerical example", 2)
    add_body(doc, "Suppose the normalised values in one window are detection 0.8, throughput 0.5, RMSE 0.2, delay 0.1, and loss 0.0. Then J = 3(0.8) + 3(0.5) − 2(0.2) − 0.1 − 0 = 3.4. If the normalised energy penalty is 0.3, the reward is 3.1. This example teaches the calculation only; it is not a measured result.")
    add_callout(doc, "Reward-design danger", "A large positive reward does not automatically mean the physical system is good. Always report the original KPIs beside reward, inspect actions, and test whether the learned policy exploits a simulator shortcut.", AMBER)

    add_heading(doc, "10. Communication KPIs are measured per window", 1)
    add_body(doc, "FlowMonitor counters increase throughout a run. If cumulative totals are read at 10, 20, and 30 seconds, the 30-second value includes earlier packets and is not the consequence of only the latest action. The implementation stores the previous counters and subtracts them. The difference belongs to the newest 10-second window.")
    add_table(doc, ["KPI", "Plain meaning", "Window calculation"], [
        ["Throughput", "Useful received data rate", "New received bytes × 8 / window duration"],
        ["Delay", "Average time for a received packet", "New delay sum / new received packets"],
        ["Loss", "Percentage transmitted but not received", "New lost packets / new transmitted packets"],
        ["Detection", "Fraction of sensing opportunities detected", "Fusion receiver’s newest window"],
        ["RMSE", "Typical localisation error", "Root mean-square position error in newest window"],
        ["Energy", "Energy used by the fleet", "Current cumulative energy − previous cumulative energy"],
    ])

    add_heading(doc, "11. The message handshake", 1)
    add_body(doc, "The simulator and Python communicate using newline-separated JSON over a local TCP socket. TCP provides ordered delivery. IDs make the semantic order explicit. Python refuses a wrong episode ID, a skipped step, a malformed message, or an observation with a size other than 50.")
    add_code(doc, '{"type":"transition", "episode_id":4, "step_id":0,\n "reward":1.25, "terminated":false, "truncated":false,\n "observation":[50 numbers], "pdet":..., "rmse_m":...}\n\n{"episode_id":4, "step_id":0, "action":9}')
    add_body(doc, "The simulator blocks after sending a non-final state. It advances only after receiving the matching action. This guarantees that one action corresponds to one following window.")

    add_heading(doc, "12. How Double DQN learns", 1)
    add_heading(doc, "12.1 The Q-network", 2)
    add_body(doc, "The network takes 50 state numbers. They pass through 256 neurons, a ReLU activation, another 256 neurons and ReLU, and finally 25 outputs. Output 0 estimates the long-term value of action 0, output 1 the value of action 1, and so on. The greedy action is the output with the largest value.")
    add_heading(doc, "12.2 The Bellman idea", 2)
    add_body(doc, "A transition teaches the network that an action’s value should be close to the reward received now plus the discounted value available next. The discount gamma is 0.99, meaning future reward matters strongly but is slightly discounted.")
    add_code(doc, "target = reward + gamma × (1 − done) × next_Q_value")
    add_heading(doc, "12.3 Why Double DQN uses two networks", 2)
    add_body(doc, "Ordinary DQN can overestimate values because one noisy network both picks and evaluates the largest next value. Double DQN uses the online network to choose the next action and the target network to value it. The target network is a delayed copy, updated every 1,000 environment steps.")
    add_heading(doc, "12.4 Why replay memory exists", 2)
    add_body(doc, "Consecutive simulator steps are strongly related. Training only on the newest transition can make learning unstable. Replay memory stores up to 100,000 transitions and samples mixed mini-batches of 128. Old and new experiences are therefore learned together.")
    add_heading(doc, "12.5 Exploration versus exploitation", 2)
    add_body(doc, "At the beginning epsilon is 1.0, so actions are random. It decreases towards 0.05 over 30,000 steps. Random exploration discovers consequences; greedy exploitation uses the best current estimate. Evaluation always turns exploration off.")
    add_heading(doc, "12.6 Loss, optimiser, and gradient clipping", 2)
    add_body(doc, "Huber loss is less sensitive to a few very large errors than squared error. AdamW updates the weights with learning rate 0.0003. Gradient norm clipping at 10 prevents an unusually large batch from causing an extreme update.")

    add_heading(doc, "13. Training, validation, and testing", 1)
    doc.add_picture(str(split), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_body(doc, "Training data change the neural-network weights. Validation data choose hyperparameters or a checkpoint. Test data are used once for the final estimate. Looking at test results and then choosing a different model leaks test information and makes the reported performance optimistic.")
    add_table(doc, ["Split", "Can update weights?", "Can choose model?", "Can make final claim?"], [
        ["Training", "Yes", "Not by training return alone", "No"],
        ["Validation", "No", "Yes", "No"],
        ["Test", "No", "No", "Yes, after the model is frozen"],
    ])
    add_body(doc, "The example full run with agent seed 31 automatically uses simulator seeds 3,100,000 through 3,100,999. Suggested validation seeds are 7001–7010 and test seeds 9001–9010. The exact sets should be declared before final evaluation.")

    add_heading(doc, "14. What was changed in the code", 1)
    changes = [
        ["uav-isac-drl.cc", "Added TCP bridge, episode/step handshake, live actions, 50-state construction, reward, termination, per-window KPI deltas, and immutable output arguments."],
        ["uav-isac-drl.cc", "Fixed radar range to include UAV-to-ground altitude and logged actual live speed/altitude targets."],
        ["uav_isac_env.py", "Added Gymnasium reset/step/close wrapper, process launch, socket validation, spaces, timeouts, command recording, and overwrite protection."],
        ["dqn_agent.py", "Added neural networks, replay buffer, Double-DQN targets, epsilon policy, Huber optimisation, gradient clipping, target updates, save/load."],
        ["mock_uav_env.py", "Added a tiny predictable environment for quick learning tests without ns-3."],
        ["tests/test_dqn.py", "Verifies Bellman target, replay shapes, and learning of a known preferred action."],
        ["tests/test_protocol.py", "Verifies message identity/action pairing and rejects the wrong state size."],
        ["tests/test_ns3_integration.py", "Verifies live sensing, speed, altitude target, and physical altitude changes."],
        ["offline_rf_benchmark.py", "Preserves and correctly names the released Random Forest method as a baseline."],
        ["generate_rf_dataset.py", "Writes deterministic configurations and exact commands for all newly generated RF runs."],
        ["select_checkpoint.py", "Evaluates candidate checkpoints only on declared validation seeds."],
        ["evaluate_policies.py", "Runs matched-seed baselines/DQN and writes raw results, method summaries, and paired differences."],
    ]
    add_table(doc, ["File", "Purpose/change"], changes)

    add_heading(doc, "15. What the tests proved", 1)
    add_table(doc, ["Test", "Result", "What it means", "What it does not mean"], [
        ["Build", "Passed", "The C++ target compiles", "The model is scientifically accurate"],
        ["Bellman target", "Passed", "Double-DQN online-selection/target-evaluation formula works", "Training will converge in ns-3"],
        ["Replay buffer", "Passed", "Transitions and batch shapes work", "Stored data are sufficient"],
        ["Mock learning", "Passed", "The code learned known action 7", "It learned good UAV control"],
        ["Protocol tests", "Passed", "IDs and 50-state messages are enforced", "TCP can never fail in deployment"],
        ["ns-3 integration", "Passed in 66.9 s", "Live controls and physical altitude changed", "The policy is better than baselines"],
        ["Three-episode smoke", "Completed", "Training/save/validate/freeze/test pipeline runs", "A DRL performance improvement"],
    ])

    add_heading(doc, "16. The measured smoke result—and how to read it", 1)
    add_table(doc, ["Method", "Return", "Detection", "RMSE (m)", "Throughput (Mbps)", "Delay (ms)", "Final-window energy (J)"], [
        ["Static", "2.6033", "0.5609", "9.4829", "0.0009184", "11.5077", "11239.26"],
        ["Random", "2.7903", "0.5808", "8.9565", "0.0009600", "11.3962", "11239.32"],
        ["Rule", "2.7215", "0.5665", "9.1766", "0.0009440", "11.5781", "11239.30"],
        ["Double DQN", "2.5076", "0.5655", "9.4867", "0.0008800", "11.9263", "11239.20"],
    ])
    add_body(doc, "This test used only seed 9002 and two action steps per policy. The DQN was trained for three 30-second episodes—six transitions total. Its epsilon was still approximately 1.0. All three validation checkpoints tied. The DQN return was lower than static in this seed. That is not a failure of the pipeline; it is the expected behaviour of an effectively untrained network.")
    add_callout(doc, "Correct sentence", "‘The smoke experiment verified end-to-end execution but was not statistically powered to compare policy performance.’", GREEN)
    add_callout(doc, "Incorrect sentence", "‘Results prove that DRL significantly improves detection and energy efficiency.’ The available data do not support this statement.", RED)

    add_heading(doc, "17. Commands: from setup to final comparison", 1)
    add_heading(doc, "17.1 Install Python packages and build", 2)
    add_code(doc, ".venv/bin/python -m pip install -r scratch/uav-isac-drl/requirements.txt\n./ns3 build uav-isac-drl -j 4")
    add_heading(doc, "17.2 Run fast unit/protocol tests", 2)
    add_code(doc, ".venv/bin/python -m unittest discover \\\n+  -s scratch/uav-isac-drl/tests -p 'test_*.py' -v")
    add_heading(doc, "17.3 Run the slow live ns-3 integration test", 2)
    add_code(doc, "RUN_NS3_INTEGRATION=1 .venv/bin/python -m unittest \\\n+  scratch/uav-isac-drl/tests/test_ns3_integration.py -v")
    add_heading(doc, "17.4 Run a random-policy smoke episode", 2)
    add_code(doc, ".venv/bin/python scratch/uav-isac-drl/random_policy_smoke.py \\\n+  --episodes 1 --seed 1001 --simulation-time 40 \\\n+  --output results/uav-isac-drl/random-smoke-new")
    add_heading(doc, "17.5 Verify DQN quickly in the mock world", 2)
    add_code(doc, ".venv/bin/python scratch/uav-isac-drl/train_dqn.py \\\n+  --environment mock --episodes 100 --learning-starts 64 --batch-size 32 \\\n+  --output results/uav-isac-drl/mock-training-new")
    add_heading(doc, "17.6 Generate a new RF baseline dataset", 2)
    add_code(doc, ".venv/bin/python scratch/uav-isac-drl/generate_rf_dataset.py \\\n+  --runs 300 --sampling-seed 20260818 --simulation-time 300 \\\n+  --output results/uav-isac-drl/rf-baseline-300")
    add_body(doc, "Use --dry-run first to inspect configurations and commands. The script refuses to append into an existing run_summary.csv because that could duplicate runs.")
    add_heading(doc, "17.7 Fit the RF benchmark", 2)
    add_code(doc, ".venv/bin/python scratch/uav-isac-drl/offline_rf_benchmark.py \\\n+  --csv results/uav-isac-drl/rf-baseline-300/run_summary.csv \\\n+  --output-dir results/uav-isac-drl/rf-model")
    add_heading(doc, "17.8 Train one full DQN candidate", 2)
    add_code(doc, ".venv/bin/python scratch/uav-isac-drl/train_dqn.py \\\n+  --environment ns3 --episodes 1000 --seed 31 --simulation-time 300 \\\n+  --learning-starts 2000 --batch-size 128 --epsilon-decay-steps 30000 \\\n+  --target-update-interval 1000 --checkpoint-every 25 \\\n+  --output results/uav-isac-drl/full/agent-seed-31")
    add_heading(doc, "17.9 Select a checkpoint using validation seeds only", 2)
    add_code(doc, ".venv/bin/python scratch/uav-isac-drl/select_checkpoint.py \\\n+  results/uav-isac-drl/full/agent-seed-31/checkpoint-episode-*.pt \\\n+  --validation-seeds 7001,7002,7003,7004,7005,7006,7007,7008,7009,7010 \\\n+  --simulation-time 300 \\\n+  --output results/uav-isac-drl/full/validation/selection.csv")
    add_heading(doc, "17.10 Freeze and compare on unseen matched seeds", 2)
    add_code(doc, ".venv/bin/python scratch/uav-isac-drl/evaluate_policies.py \\\n+  --methods static,rf-static,random,rule,dqn \\\n+  --checkpoint <path copied from selection.json> \\\n+  --rf-candidate results/uav-isac-drl/rf-model/offline_suggested_candidates.csv \\\n+  --seeds 9001,9002,9003,9004,9005,9006,9007,9008,9009,9010 \\\n+  --simulation-time 300 \\\n+  --output results/uav-isac-drl/final/comparison.csv")
    add_callout(doc, "Runtime reality", "The current debug ns-3 build is slow. The full 1,000-episode campaign can take many hours or days. Run a timed episode first, use a managed batch job, preserve logs, and do not substitute a tiny smoke run in the paper.", AMBER)

    add_heading(doc, "18. What the final paper must report", 1)
    add_bullets(doc, [
        "Exact code revision, dependency versions, machine information, and complete commands.",
        "The full state, action, reward, termination, and truncation definitions.",
        "Training episodes, decisions, replay settings, learning-start threshold, epsilon schedule, and target-update interval.",
        "Separate and predeclared training, validation, and test seed sets.",
        "At least static, random, rule-based, released RF, and frozen-DQN baselines.",
        "Mean, standard deviation, 95% confidence interval, and paired per-seed differences for every KPI.",
        "Learning curves across independent agent seeds, not a single lucky curve.",
        "Action frequencies, action saturation/bound hits, and manual trajectory inspection.",
        "Ablations: no energy term, no communication state, no mobility actions, and possibly different control intervals.",
        "Limitations: simulator abstraction, partial observation, lack of field trials, and O-RAN deployment remaining future work.",
    ])
    add_heading(doc, "18.1 Claims that are currently safe", 2)
    add_bullets(doc, [
        "A genuine sequential Double-DQN environment has been implemented.",
        "The state/action handshake and Bellman/replay logic have automated tests.",
        "Sensing, speed, and altitude actions alter live simulation state.",
        "Communication KPIs are computed from window counter differences.",
        "The complete train–validate–freeze–test pipeline executes.",
    ])
    add_heading(doc, "18.2 Claims that are not yet safe", 2)
    add_bullets(doc, [
        "DRL significantly outperforms RF or rule-based control.",
        "DRL improves detection while reducing energy.",
        "The method generalises to farms or seeds not represented by the final test campaign.",
        "The policy is ready for deployment in an O-RAN xApp or on real UAVs.",
        "The original 300-plus thesis runs have been exactly reproduced.",
    ])

    add_heading(doc, "19. Dataset recovery: what we know and what we do not", 1)
    add_body(doc, "The public repository was audited at commit 8f5ce9977b37533e4fabffcc505751cb52998713. It contains one commit, no tags, and no CSV, pickle model, experiment script, or command manifest. The supplied thesis PDF has no embedded files. Searches of the available workspace and Downloads did not find the original run summaries or RF model.")
    add_body(doc, "Therefore the exact 300-plus dataset and commands cannot honestly be recovered. The new generator reconstructs a deterministic experiment from the released bounds and records all commands, but those new runs are not the missing original thesis runs. The released RF code uses 28 features; the thesis’s count of 30 also included UE and gNB transmit power, which the released RF feature list did not use.")

    add_heading(doc, "20. Troubleshooting", 1)
    add_table(doc, ["Symptom", "Likely meaning", "What to check"], [
        ["ns-3 exits before handshake", "Build/runtime error", "Open that episode’s ns3.log and command.json"],
        ["Connection timeout", "Simulator did not open TCP port", "Build target, port availability, ns3.log"],
        ["Wrong observation size", "C++ and Python state definitions differ", "Count state insertions; keep exactly 50"],
        ["Non-consecutive step", "Protocol mismatch/stale message", "Episode ID, step ID, only one controller"],
        ["FileExistsError", "A run path already contains evidence", "Choose a new output directory; do not overwrite"],
        ["No learning loss", "Replay not ready", "learning_starts, batch size, number of transitions"],
        ["Epsilon stays near 1", "Too few steps", "Compare environment steps with epsilon_decay_steps"],
        ["Policy repeats one action", "Could be untrained, collapsed, or optimal", "Validation returns, Q-values, action counts, bounds"],
        ["Altitude changes but detection does not", "Physics or time horizon issue", "Confirm corrected slant range and enough follow-up windows"],
        ["Result changes on rerun", "Randomness/version difference", "Seed, command.json, commit, dependencies, CPU/GPU"],
    ])

    add_heading(doc, "21. Frequently asked questions", 1)
    faqs = [
        ("Is Random Forest a form of DRL?", "No. It is supervised learning here. It predicts a final score from a fixed configuration."),
        ("Why not use continuous actions?", "DQN is designed for discrete actions. The 25-action set makes the first implementation testable. Continuous control could later use SAC or TD3, but that would be a new study."),
        ("Why does the agent see herd centroid rather than all cows?", "Fifty observations keep the state compact and independent of herd size. The trade-off is loss of detailed spatial information."),
        ("Why is the first observation at 10 seconds?", "A KPI window needs elapsed time and packets. Acting at time zero would use empty measurements."),
        ("Why separate terminated and truncated?", "Battery exhaustion is a task terminal condition. Reaching the configured time limit is a truncation imposed by the experiment horizon."),
        ("Why does the static policy still use the RL bridge?", "It sends action 0 through the same timing and measurement path, so comparisons differ by policy rather than instrumentation."),
        ("What does equal-seed mean?", "Every method faces the same simulator RNG run. Paired differences then remove some scenario-to-scenario variation."),
        ("Can the smoke values go in the final Results section?", "Only as labelled implementation verification. They cannot support a performance or significance claim."),
        ("Is the trained policy already an xApp?", "No. The simulator retains an asynchronous xApp bridge, but safe deployment requires packaging, latency tests, policy constraints, monitoring, and fallback behaviour."),
        ("What is the next scientific step?", "Run the predeclared full campaign, select only on validation seeds, freeze the model, and perform the matched multi-seed test once."),
    ]
    for question, answer in faqs:
        add_heading(doc, question, 2)
        add_body(doc, answer)

    add_heading(doc, "22. Final completion checklist", 1)
    checklist = [
        "[ ] Commit or archive the exact code revision used for experiments.",
        "[ ] Record dependency and ns-3/NR versions.",
        "[ ] Declare training, validation, test, and agent seeds before the final run.",
        "[ ] Generate the RF dataset and preserve configurations.json/commands.json.",
        "[ ] Train multiple DQN candidates and independent agent seeds.",
        "[ ] Select hyperparameters/checkpoint from validation results only.",
        "[ ] Freeze the selected model and checksum the checkpoint.",
        "[ ] Run every baseline on identical unseen seeds.",
        "[ ] Generate raw, summary, and paired-vs-static CSV files.",
        "[ ] Plot learning curves, KPI distributions, action use, and energy trade-offs.",
        "[ ] Replace the red result placeholder in the LaTeX paper.",
        "[ ] Remove all wording that calls RF ‘offline DRL.’",
        "[ ] Have a second person reproduce at least one episode from command.json.",
        "[ ] Keep the limitations and distinguish simulation from field evidence.",
    ]
    for item in checklist:
        doc.add_paragraph(item)

    add_heading(doc, "Appendix A. Important output files", 1)
    add_table(doc, ["Output", "Contains", "Use"], [
        ["command.json", "Exact process argument vector", "Re-run or audit one episode"],
        ["control_kpi_log.csv", "One row per control window", "Time series, action-effect inspection"],
        ["run_summary.csv", "Final whole-episode metrics and parameters", "Episode-level comparison"],
        ["ns3.log", "Simulator output and live action traces", "Debugging and manual verification"],
        ["training.csv", "Episode returns, epsilon, loss, steps, action counts", "Learning diagnostics"],
        ["checkpoint-episode-*.pt", "Frozen candidate neural-network state", "Validation selection"],
        ["selection.csv/json", "Validation returns and selected path", "Proof that test seeds were not used"],
        ["comparison.csv", "One row per method and test seed", "Raw final evidence"],
        ["comparison-summary.csv", "Means, SDs, 95% intervals", "Paper summary table"],
        ["comparison-paired-vs-static.csv", "Matched-seed differences", "Fair policy comparison"],
    ])

    add_heading(doc, "Appendix B. Source map", 1)
    add_body(doc, "Implementation directory:")
    add_code(doc, str(ROOT / "scratch/uav-isac-drl"))
    add_body(doc, "Verified post-fix smoke artifacts:")
    add_code(doc, str(ROOT / "results/uav-isac-drl/ns3-smoke-training-fixed") + "\n" +
                  str(ROOT / "results/uav-isac-drl/ns3-smoke-validation-fixed") + "\n" +
                  str(ROOT / "results/uav-isac-drl/ns3-smoke-comparison-fixed"))
    add_body(doc, "The directories without the -fixed suffix are explicitly invalidated because they predate the altitude/radar-range correction.")

    add_heading(doc, "Appendix C. Further reading", 1)
    add_bullets(doc, [
        "R. Sutton and A. Barto, Reinforcement Learning: An Introduction, 2nd ed., MIT Press, 2018.",
        "V. Mnih et al., ‘Human-level control through deep reinforcement learning,’ Nature, 2015. https://doi.org/10.1038/nature14236",
        "H. van Hasselt, A. Guez, and D. Silver, ‘Deep Reinforcement Learning with Double Q-Learning,’ AAAI, 2016. https://doi.org/10.1609/aaai.v30i1.10295",
        "M. Towers et al., ‘Gymnasium: A Standard Interface for Reinforcement Learning Environments,’ 2024. https://arxiv.org/abs/2407.17032",
    ])

    # Start a final short revision record on its own page.
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Revision record", 1)
    add_table(doc, ["Date", "Version", "Scope"], [["18 August 2026", "1.0", "Initial beginner guide based on verified post-fix code and smoke artifacts"]])
    add_callout(doc, "Document status", "The explanation and implementation evidence are complete. Scientific performance results remain pending the full multi-seed campaign.", AMBER)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
